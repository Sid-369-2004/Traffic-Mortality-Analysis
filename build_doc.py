import docx

def update_report():
    doc = docx.Document('AD Lab - Report Template.docx')
    
    for para in doc.paragraphs:
        txt = para.text.strip()
        
        # Chapter 1
        if "A brief introduction to the topic assigned" in txt:
            para.text = "The Traffic Mortality Analysis system is an advanced AI-powered diagnostic tool designed to predict traffic accident severity. By analyzing a massive dataset of real-world traffic collisions, the system evaluates 11 critical environmental and infrastructure dimensions—including coordinate geography, meteorological pressure, humidity, wind speed, and junction proximity. The primary objective is to categorize incidents as either 'Non-Severe' or 'Severe/Fatal'. The problem is resolved using Python's full-stack capabilities, integrated with a Flask web architecture for dynamic deployment. The intelligence engine actively compares three supervised machine learning algorithms: the Multi-Layer Perceptron (MLP) Neural Network, Decision Trees, and Support Vector Machines (SVM)."
        
        # Chapter 2
        elif "This section should describe the issues determined" in txt:
            para.text = "Problem Statement:\nTraffic accidents remain a leading cause of mortality and urban infrastructure degradation. Emergency responders and civil engineers struggle to preemptively identify which accident profiles are statistically likely to involve severe casualties versus minor property damage. This creates massive inefficiencies in emergency routing and urban planning."
        elif "The objectives are determined by how" in txt:
            para.text = "Objectives:\n1. To implement strict data integrity by purging incomplete records (NaN values) rather than utilizing statistical imputation, ensuring the model only learns from 100% verified accident reports.\n2. To utilize TruncatedSVD dimensionality reduction to mathematically compress 11 environmental factors into 8 components, optimizing the system for real-time web deployment.\n3. To evaluate and declare a primary inference engine by actively comparing three distinct machine learning algorithms against unseen testing data."
        
        # Chapter 3
        elif "Chapter 3" in txt:
            # We will just append the packages safely in the next blank paragraphs
            pass
            
        elif txt == "Python Packages Used Details":
            # The template leaves this blank below the header, we will add text to the next empty paragraph
            pass

        # Chapter 4
        elif "The detailed source code is to be provided here" in txt:
            para.text = "The core intelligence is split between the Machine Learning pipeline (train.py) and the Web Architecture (app.py). Both rely on the TruncatedSVD algorithms to feed data to the MLP Neural Network.\n\n[STUDENT ACTION REQUIRED: Paste your Vercel Deployment Link here, and insert screenshots of your python code below to satisfy the rubric.]"
            
        # Chapter 5
        elif "The obtained results in the form of graphs" in txt:
            para.text = "The models were trained on an 80/20 train/test split of fully sanitized data (Strict NaN Omission).\n- The MLP Neural Network achieved the highest balanced performance across Accuracy, Precision, and F1-Score, making it the primary analytical engine.\n- Decision Trees showed high baseline accuracy but struggled with complex recall patterns.\n- Support Vector Machines (SVM) provided a rigid mathematical boundary but were computationally heavier.\n\n[STUDENT ACTION REQUIRED: Open your Vercel /analytics page and copy-paste the two beautiful Glassmorphic charts into this section.]"
            
        # Chapter 6
        elif "Conclude the report using the findings" in txt:
            para.text = "The deployment of the Traffic Mortality Analysis system successfully demonstrates that environmental and infrastructural metadata can mathematically predict the severity of traffic collisions. By enforcing strict data sanitation and utilizing an MLP Neural Network optimized via TruncatedSVD, the system achieves rapid inference speeds suitable for real-world deployment on Edge networks (Vercel). The project proves that integrating advanced Python AI algorithms into a full-stack web application yields a professional, government-grade utility for emergency risk assessment."

    # For Chapter 3 which had no placeholder text, we manually find the Chapter 3 header and insert
    for i, para in enumerate(doc.paragraphs):
        if "Python Packages Used Details" in para.text:
            packages = [
                "1. Pandas: Utilized for massive CSV dataset manipulation, strict NaN row dropping, and dataframe structuring.",
                "2. Scikit-Learn: The core mathematical engine used for training the MLP Neural Network, Decision Tree, and SVM. Also provided TruncatedSVD for feature reduction.",
                "3. Flask: The WSGI web framework utilized to construct the multi-page web application and REST API endpoints.",
                "4. Joblib: Used to serialize (pickle) the trained AI models and scalers into binary files for instantaneous web loading.",
                "5. Numpy: Handled the multi-dimensional array mathematics required by the neural network during inference."
            ]
            for pkg in reversed(packages):
                doc.paragraphs[i].insert_paragraph_before(pkg)

    doc.save('Final_AD_Lab_Report.docx')
    print("Successfully generated Final_AD_Lab_Report.docx")

if __name__ == "__main__":
    update_report()
